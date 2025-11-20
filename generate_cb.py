#!/usr/bin/env python3
"""
generate_cb_full.py — One-click exporter for your whole SMW-v1 codebase into a
single Word document.

Usage:
  python generate_cb_full.py

Behavior:
  • Scans the CURRENT WORKING DIRECTORY recursively (run it from SMW-v1 root).
  • Skips common junk/binary folders and file types (.git, node_modules, images, etc.).
  • Includes .env and other text config files.
  • Writes a single Word document 'codebase_full.docx' with file-by-file content,
    using monospace formatting.

Tip:
  If Word says permission denied, close codebase_full.docx and run again.
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# ====== SETTINGS ======
ROOT = Path(__file__).resolve().parent
OUTPUT_FILE = ROOT / "codebase_full.docx"

# Directories we completely skip
EXCLUDED_DIRS = {
    ".git",
    ".idea",
    ".vscode",
    "__pycache__",
    "node_modules",
    ".next",
    ".turbo",
    ".venv",
    "venv",
    ".pytest_cache",
    ".mypy_cache",
    "dist",
    "build",
}

# File extensions to SKIP (binary/irrelevant for code export)
SKIP_EXTENSIONS = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".ico",
    ".svg",
    ".webp",
    ".pdf",
    ".ttf",
    ".woff",
    ".woff2",
    ".eot",
    ".otf",
    ".mp3",
    ".mp4",
    ".mov",
    ".avi",
    ".zip",
    ".tar",
    ".gz",
    ".bz2",
    ".7z",
    ".exe",
    ".dll",
    ".so",
    ".dylib",
    ".obj",
    ".pyc",
    ".pyo",
    ".class",
    ".db",
    ".sqlite3",
    ".docx",   # avoid including our own output docx or others
}


def setup_code_style(doc: Document):
    """
    Ensure there is a 'Code' character style with monospace font.
    No low-level XML tricks, just standard python-docx API.
    """
    styles = doc.styles
    try:
        code_style = styles["Code"]
    except KeyError:
        code_style = styles.add_style("Code", WD_STYLE_TYPE.CHARACTER)

    font = code_style.font
    font.name = "Consolas"
    font.size = Pt(9)
    return code_style


def is_binary_file(path: Path) -> bool:
    """
    Simple heuristic: skip by extension using SKIP_EXTENSIONS.
    """
    return path.suffix.lower() in SKIP_EXTENSIONS


def iter_files(root: Path):
    """
    Yield all files under root, respecting EXCLUDED_DIRS.
    """
    for dirpath, dirnames, filenames in os.walk(root):
        # remove excluded dirs in-place so os.walk won't descend into them
        dirnames[:] = [d for d in dirnames if d not in EXCLUDED_DIRS]

        for filename in filenames:
            full_path = Path(dirpath) / filename
            # skip the output file itself, if present in tree
            if full_path.resolve() == OUTPUT_FILE.resolve():
                continue
            if is_binary_file(full_path):
                continue
            yield full_path


def main():
    print(f"Project root: {ROOT}")
    print(f"Output file : {OUTPUT_FILE}")
    print("Building Word document...")

    doc = Document()
    doc.core_properties.title = "SMW-v1 Codebase"
    doc.add_heading("SMW-v1 Codebase Export", level=0)
    doc.add_paragraph("Generated from project root: {}".format(ROOT))

    code_style = setup_code_style(doc)

    for path in sorted(iter_files(ROOT)):
        rel_path = path.relative_to(ROOT).as_posix()
        print(f"Adding: {rel_path}")

        # Add a heading for this file
        doc.add_heading(rel_path, level=2)

        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            # If something goes wrong, note it in the doc
            p = doc.add_paragraph()
            run = p.add_run(f"[Error reading file: {e}]")
            run.italic = True
            continue

        if not text:
            p = doc.add_paragraph()
            run = p.add_run("[Empty file]")
            run.italic = True
            continue

        # Add file content as a "code" run
        # (newline characters in `text` become Word line breaks)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.style = code_style

        # Spacer between files
        doc.add_paragraph()

    doc.save(str(OUTPUT_FILE))
    print("Done. Saved:", OUTPUT_FILE)


if __name__ == "__main__":
    main()
